import csv
from docx import Document
from docx.shared import RGBColor, Inches, Pt
import datetime
import os

# Build usable CSV file
exec(open('crosstabcleanup.py').read())
# Sort the data
exec(open('sortbyday.py').read())
# Create capacity dictionaries
exec(open('determinecapacity.py').read())
# Populate other necessary variables
exec(open('extravars.py').read())

# create the document
document = Document()
sections = document.sections
for section in sections:
    section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = Inches(.5), Inches(.5), Inches(.75), Inches(.75)
document.add_heading('VIRTUAL TRAINING COACH | CLASS ASSIGNMENTS BY ROOM', level=1)
warning = document.add_paragraph().add_run('Enrollment numbers are subject to change.')
font = warning.font
font.color.rgb = RGBColor(255, 0, 0)
footer = section.footer
paragraph = footer.paragraphs[0]
today = datetime.datetime.today().strftime('%m.%d.%Y')
paragraph.text = f"VIRTUAL TRAINING PROGRAM | COACH CLASS AND ROOM ASSIGNMENTS | REVISED {today}"


def createtable(day, dow, date, enrolled, coachone, coachtwo, coachthree):
    # Create capacity dictionaries
    currentdaycap = dict(roomcapacity)
    coachesusedprimary = []
    coachesusedsecondary = []
    coachesusedthird = []
    document.add_heading(f'{dow.upper()}, {date}', level=2)
    table = document.add_table(rows=1, cols=6)
    table.style = document.styles['Medium Shading 1 Accent 5']
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Location-Room"
    hdr_cells[1].text = "Coach"
    hdr_cells[2].text = "Class"
    hdr_cells[3].text = "Start Time"
    hdr_cells[4].text = "End Time"
    hdr_cells[5].text = "Students"
    hdr_cells[2].width = Inches(3)
    currentval = ""
    for site in ["BTC", "LTC", "OTC", "SFTC", "SRTC", "STC", "SMPS", "MMC", "RS3TC", "GWO"]:
        for info in day:
            if site == info[2]:
                while enrolled[info[2]] > 0:
                    if currentdaycap[primaryrooms[info[2]]]-int(info[-1]) >= 0:
                        currentdaycap[primaryrooms[info[2]]] -= int(info[-1])
                        enrolled[info[2]] -= int(info[-1])
                        try:
                            row_cells = table.add_row().cells
                            row_cells[2].text = info[3]
                            row_cells[3].text = info[4]
                            row_cells[4].text = info[5]
                            row_cells[5].text = info[-1]
                            if currentval == coachone[info[2]]:
                                row_cells[0].text = ""
                                row_cells[1].text = ""
                            else:
                                row_cells[0].text = primaryrooms[info[2]]
                                row_cells[1].text = coachone[info[2]]
                                currentval = coachone[info[2]]
                                coachesusedprimary.append(coachone[info[2]])
                        except KeyError:
                            print("Is your Master Capacity spreadsheet updated? Showing enrollment for a site that isn't "
                                  "listed.")
                            print(day)
                            pass
                        break
                    elif currentdaycap[secondaryrooms[info[2]]]-int(info[-1]) >= 0:
                        currentdaycap[secondaryrooms[info[2]]] -= int(info[-1])
                        enrolled[info[2]] -= int(info[-1])
                        try:
                            row_cells = table.add_row().cells
                            row_cells[2].text = info[3]
                            row_cells[3].text = info[4]
                            row_cells[4].text = info[5]
                            row_cells[5].text = info[-1]
                            if currentval == coachtwo[info[2]]:
                                row_cells[0].text = ""
                                row_cells[1].text = ""
                            else:
                                row_cells[0].text = secondaryrooms[info[2]]
                                row_cells[1].text = coachtwo[info[2]]
                                currentval = coachtwo[info[2]]
                                coachesusedsecondary.append(coachtwo[info[2]])
                        except KeyError:
                            print("Is your Master Capacity spreadsheet updated? Showing enrollment for a site that isn't "
                                  "listed.")
                            pass
                        break
                    elif currentdaycap[thirdrooms[info[2]]]-int(info[-1]) >= 0:
                        currentdaycap[thirdrooms[info[2]]] -= int(info[-1])
                        enrolled[info[2]] -= int(info[-1])
                        try:
                            row_cells = table.add_row().cells
                            row_cells[2].text = info[3]
                            row_cells[3].text = info[4]
                            row_cells[4].text = info[5]
                            row_cells[5].text = info[-1]
                            if currentval == coachthree[info[2]]:
                                row_cells[0].text = ""
                                row_cells[1].text = ""
                            else:
                                row_cells[0].text = thirdrooms[info[2]]
                                row_cells[1].text = coachthree[info[2]]
                                currentval = coachthree[info[2]]
                                coachesusedthird.append(coachthree[info[2]])
                        except KeyError:
                            print("Is your Master Capacity spreadsheet updated? Showing enrollment for a site that isn't listed.")
                            print(day)
                            pass
                    else:
                        enrolled[info[2]] -= int(info[-1])
            else:
                pass

        def primarynotsecheduled():
            cell = row_cells[2]
            pt = cell.paragraphs[0]
            t = pt.text = ''
            pt.add_run("For current enrollments, check ").font.color.rgb = RGBColor(255, 0, 0)
            run = pt.add_run("Healthstream and the Site Report")
            pt.add_run(" on the preceding business day.").font.color.rgb = RGBColor(255, 0, 0)
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)
            run.underline = True
            run.space_before = Pt(0)

        def secondarynotscheduled():
            cell = row_cells[2]
            pt = cell.paragraphs[0]
            t = pt.text = ''
            pt.add_run("Check ").font.color.rgb = RGBColor(255, 0, 0)
            run = pt.add_run("Healthstream and the Site Report")
            pt.add_run(" on the preceding business day. if enrollment is ").font.color.rgb = RGBColor(255, 0, 0)
            runtwo = pt.add_run(f"{roomcapacity[primaryrooms[coach]]} or less at one time,")
            pt.add_run(" shift is relieved.").font.color.rgb = RGBColor(255, 0, 0)
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)
            run.underline = True
            font = runtwo.font
            font.color.rgb = RGBColor(255, 0, 0)
            runtwo.underline = True
            run.space_before = Pt(0)

        def thirdnotsecheduled():
            cell = row_cells[2]
            pt = cell.paragraphs[0]
            t = pt.text = ''
            pt.add_run("Check ").font.color.rgb = RGBColor(255, 0, 0)
            run = pt.add_run("Healthstream and the Site Report")
            pt.add_run(" on the preceding business day. if enrollment is ").font.color.rgb = RGBColor(255, 0, 0)
            runtwo = pt.add_run(f"{roomcapacity[primaryrooms[coach]] + roomcapacity[secondaryrooms[coach]]}"
                                f" or less at one time,")
            pt.add_run(" shift is relieved.").font.color.rgb = RGBColor(255, 0, 0)
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)
            run.underline = True
            font = runtwo.font
            font.color.rgb = RGBColor(255, 0, 0)
            runtwo.underline = True
            run.space_before = Pt(0)

        for coach in coachone:
            if coachone[coach] not in coachesusedprimary and site == coach:
                row_cells = table.add_row().cells
                row_cells[0].text = primaryrooms[coach]
                row_cells[1].text = coachone[coach]
                row_cells[2].merge(row_cells[3]).merge(row_cells[4]).merge(row_cells[5])
                primarynotsecheduled()
        for coach in coachtwo:
            if coachtwo[coach] not in coachesusedsecondary and site == coach:
                row_cells = table.add_row().cells
                row_cells[0].text = secondaryrooms[coach]
                row_cells[1].text = coachtwo[coach]
                row_cells[2].merge(row_cells[3]).merge(row_cells[4]).merge(row_cells[5])
               secondarynotscheduled()
        for coach in coachthree:
            if coachthree[coach] not in coachesusedthird and site == coach:
                row_cells = table.add_row().cells
                row_cells[0].text = thirdrooms[coach]
                row_cells[1].text = coachthree[coach]
                row_cells[2].merge(row_cells[3]).merge(row_cells[4]).merge(row_cells[5])
                thirdnotsecheduled()


createtable(monday, "Monday", nextmon, monenrolled, primarycoachmon, secondarycoachmon, thirdcoachmon)
document.add_page_break()
createtable(tuesday, "Tuesday", nexttue, tuesenrolled, primarycoachtues, secondarycoachtues, thirdcoachtues)
document.add_page_break()
createtable(wednesday, "Wednesday", nextwed, wedsenrolled, primarycoachwed, secondarycoachwed, thirdcoachwed)
document.add_page_break()
createtable(thursday, "Thursday", nextthu, thurenrolled, primarycoachthur, secondarycoachthur, thirdcoachthur)
document.add_page_break()
createtable(friday, "Friday", nextfri, frienrolled, primarycoachfri, secondarycoachfri, thirdcoachfri)

if os.path.exists("I:/CrosstabOutputTemp.csv"):
    os.remove("I:/CrosstabOutputTemp.csv")
document.save(f'I:/{today}.VTCoach_ClassAssignmentsbyRoom.docx')
